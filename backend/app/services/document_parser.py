import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentParser:
    @staticmethod
    def parse_pdf(content: bytes) -> Optional[str]:
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            result = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(result)} characters from PDF")
            return result if result.strip() else None

        except ImportError:
            logger.error("pdfplumber not installed, trying PyPDF2")
            try:
                import PyPDF2

                text_parts = []
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                result = "\n\n".join(text_parts)
                logger.info(f"Extracted {len(result)} characters from PDF using PyPDF2")
                return result if result.strip() else None

            except Exception as e:
                logger.error(f"Failed to parse PDF with PyPDF2: {e}")
                return None

        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            return None

    @staticmethod
    def parse_docx(content: bytes) -> Optional[str]:
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            result = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(result)} characters from DOCX")
            return result if result.strip() else None

        except ImportError:
            logger.error("python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            return None

    @staticmethod
    def parse_doc(content: bytes) -> Optional[str]:
        try:
            import subprocess
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp_in:
                tmp_in.write(content)
                tmp_in_path = tmp_in.name

            tmp_out_path = tmp_in_path.replace(".doc", ".txt")

            try:
                result = subprocess.run(
                    ["antiword", tmp_in_path],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )

                if result.returncode == 0 and result.stdout.strip():
                    logger.info(
                        f"Extracted {len(result.stdout)} characters from DOC using antiword"
                    )
                    return result.stdout
                else:
                    logger.warning("antiword failed or returned empty text")
                    return None

            finally:
                if os.path.exists(tmp_in_path):
                    os.unlink(tmp_in_path)
                if os.path.exists(tmp_out_path):
                    os.unlink(tmp_out_path)

        except FileNotFoundError:
            logger.error("antiword not installed, cannot parse .doc files")
            return None
        except Exception as e:
            logger.error(f"Failed to parse DOC: {e}")
            return None

    @staticmethod
    def parse_document(content: bytes, filename: str) -> Optional[str]:
        ext = filename.lower().split(".")[-1] if "." in filename else ""

        if ext == "pdf":
            return DocumentParser.parse_pdf(content)
        elif ext == "docx":
            return DocumentParser.parse_docx(content)
        elif ext == "doc":
            return DocumentParser.parse_doc(content)
        elif ext == "txt":
            try:
                return content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    return content.decode("cp1251")
                except Exception as e:
                    logger.error(f"Failed to decode text file: {e}")
                    return None
        else:
            logger.error(f"Unsupported file type: {ext}")
            return None
